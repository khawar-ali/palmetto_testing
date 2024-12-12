import streamlit as st
from docx import Document
from spire.doc import Document as dct
from spire.doc.common import *
from langchain_openai import AzureChatOpenAI
import tempfile
import pandas as pd
import os
import re
from dotenv import load_dotenv
import openai
import base64  # Import the base64 module for encoding
from io import BytesIO  # Import BytesIO from the io module
from xml.etree import ElementTree as ET
from docx.shared import Pt
# Load environment variables from .env file
load_dotenv()

# Access the API key
#OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
#OPENAI_API_ORG_KEY = os.getenv("OPENAI_API_ORG_KEY")
#azure_api_key = os.getenv("AZURE_API_KEY")
#azure_endpoint = os.getenv("AZURE_ENDPOINT")

azure_api_key = st.secrets["AZURE_API_KEY"]
azure_endpoint = st.secrets["AZURE_ENDPOINT"]
# Set the API key for OpenAI
#openai.api_key = OPENAI_API_KEY
#openai.organization = OPENAI_API_ORG_KEY

print("Streamlit app is running.")
st.session_state.head_match= ""
if 'final_result' not in st.session_state:
    st.session_state.final_result = ""
BookTitle = ""
AuthorName = ""
SubTitle=""
Description = ""
Audience = ""
Genre = ""
Headline = ""
newliner= "\n"
OnlineBookDescription =""
LongBackCoverCopy=""
KeywordsForOnlineDiscoverability=""
AuthorBiography=""

screenOutput = ""
file_name_input = ""
Book_title_prompt= """
    [Instructions]:
    Place the book's title {BookTitle} here. Do not make this up, only copy and paste from the provided document.
"""
Author_name_prompt = """ 
    [Instructions]
    Place the author's name {AuthorName} here. Do not make this up, only copy and paste from the provided document.
"""

Book_Subtitle_prompt="""
    [Instructions]
    Place the book's subtitle {SubTitle} here. Do not make this up, only copy and paste from the provided document.
"""
Tagline_Suggestions_prompt =""" Create 2-5 tagline suggestions for the book described in {Description}. 
The taglines should not include {BookTitle} or {AuthorName}. 
Sample taglines: 
    "Nothing is more dangerous than a faerie tale."  
    "Trouble just moved in." 
    "Power is a dangerous game." 
    "Sometimes love can be arranged."
All content should be written at a reading level that is appropriate for the {Audience}. Age plays a factor in what a group of readers will be able to understand. Make sure all output ALWAYS stays at the reading level of the {Audience}. 
The tagline suggestions should use differing sentence structures and avoid sentence structure patterns.
"""

Keywords_Online_Discoverability_prompt = """ Create 15 Keywords for Online Discoverability based on the book in {Description}. 
These keywords are related to the book and will help with online discoverability. 
The 15 keywords should be separated by semicolons. For example this section should look exactly like “Greece; Greek islands; memoir; travel; self-discovery; healing; mental health; LGBTQ+; women's issues; cultural exploration; family relationships; personal growth; adventure; humor; resilience; pandemic; vacation; self-reflection; empowerment; emotional journey'' but the keywords will change and be based on the book {Description}. 
Do not include character's or people's names in the keywords.
"""

BISAC_Codes_prompt=""" Provide only 3 BISAC Codes relating to the {Genre} of the book provided in {Description} that are as specific as possible. 
NEVER make up the BISAC Codes. 
ONLY use accurate and real codes that relate to the {Genre} in the output. 
Keep this in a list format separated by semicolons. For example this section should look exactly like “JUV039140 JUVENILE FICTION ; Social Themes ; Friendship JUV039060 JUVENILE FICTION ; Social Themes ; Self-Esteem & Self-Reliance JUV002030 JUVENILE FICTION ; Animals ; Frogs & Toads'' but the codes and words will change and be relevant to the book {Genre}. 
Keep to one genre, but include 3 BISAC codes. 
ONLY use accurate and real codes in the output. 
"""
Online_Book_Description_prompt= """  Write a unique 300-word online book description using {Description} for context.
Never assume or make up the book title. Always use the book title from {BookTitle} where required. 
The output should be in paragraph format. 
Use {Headline} as the only content in the first paragraph of Online Book Description. The only text in the first paragraph is {Headline}.
Do not use wording from {Headline} afterward the first paragraph.
The output should be written at a reading level that is appropriate for {Audience}. 
Do not mention details about {Audience} in the output.
The output MUST stay as close to 300 words as possible.
Always avoid repetitive phrasing and wording.
"""


Long_Back_Cover_Copy_prompt=""" Write a 250-word summary using {Description} for context.
The output should be in paragraph format. 
The output must be different from the {OnlineBookDescription}.
Avoid exact or similar sentences and phrases found in the {OnlineBookDescription}.
All content should be written at a reading level that is appropriate for {Audience}. 
Do not mention details about {Audience} in the output.
All output MUST stay as close to 250 words as possible. 
Always avoid repetitive phrasing and wording. Always use book title from {BookTitle} where required.
"""
Short_Back_Cover_Copy_prompt=""" Write a unique 75-word summary using {Description} for context.
The output should be in paragraph format. 
The output must be different from the {OnlineBookDescription} and {LongBackCoverCopy}.
Avoid exact or similar sentences and phrases found in the {OnlineBookDescription} and {LongBackCoverCopy}.
The goal is summarizing the plot to make the book desirable to a reader in 75 words. 
All content should be written at a reading level that is appropriate for {Audience}. 
Do not mention details about {Audience} in the output.
All output MUST stay as close to 75 words as possible. 
Always avoid repetitive phrasing and wording. Always use book title from {BookTitle} where required.


"""
about_the_author_prompt=""" Write a 100 word author biography using this context {AuthorBiography}. 
The goal is to give the reader information about the author. Find this information in the document provided to you and do not make information about the author up. 
The book is called {BookTitle} and the author is {AuthorName}. 
Keep the author biography’s language in line with language and reading levels that would be appropriate for {Audience}. 
If there is more than one author, create multiple but separate biographies. Each author on a project MUST have their own 100 word biography/ description. 
All outputs Must stay as close to 100 words as possible. 
NEVER make up or create new information about the author or authors. 
Only use the provided document with information about the author to write this biography. It should NEVER include false information, degrees, or fields they are renowned in unless the information is specifically included in the provided document.
Always avoid repetitive phrasing and wording.

"""
Headline_prompt=""" From {Description}, generate a 239 character headline.
The 239 character limit includes spaces and punctuation.
NEVER include quotation marks or brackets in the output. 
Do not include {AuthorName}, {BookTitle}, or {SubTitle}. 
All content should be written at a reading level that is appropriate for {Audience}. 
Do not include any language from the prompt in the output.
Output MUST be as close to 239 characters as possible.
"""

def process_book_text(book_text):
    # Initialize variables
    book_title = author_name = book_subtitle = ""
    tagline_suggestions = keywords = bisac_codes = ""
    online_book_description = long_back_cover_copy = short_back_cover_copy = ""
    about_the_author = headline = ""

    # Extract Book Title
    print("==================================================")
    print(book_text)
    print("==================================================")
   # bt_match = re.search(r"Book Title:\n(.+)\n", book_text)
    bt_match = re.search(r"Book Title:(.+?)\n\nAuthor Name:", book_text, re.DOTALL)

    print(bt_match)
    if bt_match:
        book_title = bt_match.group(1).strip()
       # book_title = f"<i>{book_title}</i>"
        #book_title = f"*{book_title}*"
    # Extract Author Name
    #a_match = re.search(r"Author Name:\n(.+)\n", book_text)
    a_match = re.search(r"Author Name:(.+?)\n\nBook Subtitle:", book_text, re.DOTALL)

    if a_match:
        author_name = a_match.group(1).strip()

    # Extract Book Subtitle
    #match = re.search(r"Book Subtitle:\n(.+)\n", book_text)
    match = re.search(r"Book Subtitle:(.+?)\n\nTagline Suggestions:", book_text, re.DOTALL)

    if match:
        book_subtitle = match.group(1).strip()

    # Extract Tagline Suggestions
   
    #match = re.search(r"Tagline Suggestions:(.+?)(?=\w+:|$)", book_text, re.DOTALL)
   # match = re.search(r"Tagline Suggestions:(.+?)(?=\w+|$)", book_text, re.DOTALL)
    match = re.search(r"Tagline Suggestions:(.+?)\n\nKeywords for Online Discoverability:", book_text, re.DOTALL)
    if match:
        tagline_suggestions = match.group(1).strip()

    # Extract Keywords
    #match = re.search(r"Keywords for Online Discoverability:(.+?)(?=\w+:|$)", book_text, re.DOTALL)
    match = re.search(r"Keywords for Online Discoverability:(.+?)\n\nBISAC Codes:", book_text, re.DOTALL)

    if match:
        keywords = match.group(1).strip()

    # Extract BISAC Codes
   # match = re.search(r"BISAC Codes:(.+?)(?=\w+:|$)", book_text, re.DOTALL)
    match = re.search(r"BISAC Codes:(.+?)\n\nOnline Book Description:", book_text, re.DOTALL)

    if match:
        bisac_codes = match.group(1).strip()

    # Extract Online Book Description
    #match = re.search(r"Online Book Description:(.+?)(?=\w+:|$)", book_text, re.DOTALL)
    match = re.search(r"Online Book Description:(.+?)\n\nLong Back Cover Copy:", book_text, re.DOTALL)

    if match:
        online_book_description = match.group(1).strip()

    # Extract Long Back Cover Copy
    #match = re.search(r"Long Back Cover Copy:(.+?)(?=\w+:|$)", book_text, re.DOTALL)
    match = re.search(r"Long Back Cover Copy:(.+?)\n\nShort Back Cover Copy:", book_text, re.DOTALL)

    if match:
        long_back_cover_copy = match.group(1).strip()

    # Extract Short Back Cover Copy
    #match = re.search(r"Short Back Cover Copy:(.+?)(?=\w+:|$)", book_text, re.DOTALL)
    match = re.search(r"Short Back Cover Copy:(.+?)\n\nAbout the Author:", book_text, re.DOTALL)

    if match:
        short_back_cover_copy = match.group(1).strip()

    # Extract About the Author
    match = re.search(r"About the Author:(.+?)\n\nHeadline:", book_text, re.DOTALL)
    if match:
        about_the_author = match.group(1).strip()
    #match = re.search(r"Headline:(.+?)\n\n", book_text, re.DOTALL) 
    match = re.search(r"Headline:(.+)", book_text, re.DOTALL)   
    #match = re.search(r"Headline:(.+?)(?=\w+:|$)", book_text, re.DOTALL)
    if match:
        headline = match.group(1).strip()    

    # Return a dictionary with extracted variables
    #"online_book_description": headline + newliner + online_book_description,
    return {
        "book_title": book_title,
        "author_name": author_name,
        "book_subtitle": book_subtitle,
        "tagline_suggestions": tagline_suggestions,
        "keywords": keywords,
        "bisac_codes": bisac_codes,
        "online_book_description": online_book_description,
        "long_back_cover_copy": long_back_cover_copy,
        "short_back_cover_copy": short_back_cover_copy,
        "about_the_author": about_the_author,
        "headline":headline,
    }

def extract_tabular_data(document):
        # Get the first section of the document
    section = document.Sections[0]
    # Get the first table in the section
    table = section.Tables[0]
    # Create a list to store the extracted table data
    data_list = []
    columns=['Book_details','Book_details_1']
    df=pd.DataFrame(columns=columns)
    # Loop through the rows in the table
    all_data=[]
    for i in range(table.Rows.Count):
        row = table.Rows[i]
        # Loop through the cells in each row
        j=0
        omit_row=''
        if len(range(row.Cells.Count))==1:
            omit_row='omit'
        else:
            text_cell=''
            for j in range(row.Cells.Count):
                cell = row.Cells[j]
                for k in range(cell.Paragraphs.Count):
                    # Extract data from each paragraph
                    paragraph = cell.Paragraphs[k]
                    text = paragraph.Text
                    # Append the data to the list
                    data_list.append(text)
                    text_cell+=text
                
                data_list.append("\t")
                if j==0:
                  all_data.append({'Book_details': text_cell})

                else: 
                  all_data[-1]['Book_details_1']= text_cell
                text_cell=''   

            if not omit_row=='omit':
                data_list.append('\n')    

    df=pd.concat([df,pd.DataFrame(all_data)],ignore_index=True)
    return df
def read_and_update_tables(doc, variable_data):
    # Create a new document to store the updated content  # writing data in template.docx
    updated_document = doc
    # Assuming the book title is the first item in variable_data and removing double quotes
    book_title = variable_data[0].replace('"', '')  
    #book_title = book_title.lower()
    st.session_state.head_match = variable_data[10].replace('"', '')  
    for table_index, table in enumerate(doc.tables):
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # Update the second column with variable data
                if j == 1 and i < len(variable_data):
                    # Clear the existing text and formatting in the cell
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    cell_content = variable_data[i].replace('"', '')  # Remove double quotes from cell content

                    if i == 0:  # Book title in row 0, column 1
                        # Apply italic formatting to the entire book title
                        run = paragraph.add_run(cell_content)
                        run.italic = True
                    elif i == 6:  # Specific case for row 6, column 1
                        # Split the text into lines
                        lines = cell_content.split('\n')
                        print('lines',lines)
                        if lines:
                            # Apply bold formatting to the first line
                            run = paragraph.add_run(lines[0])
                            print("-------------Headline----------")
                            #print(run.text)
                            print(st.session_state.head_match)
                            if st.session_state.head_match in lines[0]:
                                print("match found")
                            #    line[0]+='.'
                                index = lines[0].find(st.session_state.head_match)
                                
                                # Apply bold formatting to the specific part of the line
                                #line[0]+='.'
                                run.text = lines[0][:index]
                                run = paragraph.add_run(lines[0][index:index+len(st.session_state.head_match)])
                                run.bold = True
                                # Add the rest of the line without bold formatting to the existing run
                                run = paragraph.add_run(lines[0][index+len(st.session_state.head_match):])
                            paragraph.add_run('\n')
                            # Add the rest of the lines normally
                            for line in lines[1:]:
                                # Check and apply italic formatting for book title within these lines
                                
                                parts = line.split(book_title)
                                paragraph.add_run('\n')
                                for k, part in enumerate(parts):
                                    
                                    paragraph.add_run(part)
                                    if k < len(parts) - 1:  # Add the book title in italic, if not the last part

                                        italic_run = paragraph.add_run(book_title.replace("'",""))
                                        italic_run.italic = True
                    else:
                        # For other cells, check for book title and apply italic formatting
                        parts = cell_content.split(book_title)
                        for k, part in enumerate(parts):
                            paragraph.add_run(part)
                            if k < len(parts) - 1:  # Add the book title in italic, if not the last part
                                italic_run = paragraph.add_run(book_title.replace("'",""))
                                italic_run.italic = True

    return updated_document


def copy_table_format(source_docx, target_docx):
    # Load the source and target documents
    source_doc = Document(source_docx)
    target_doc = Document(target_docx)
    # Assuming each document contains only one table and we are working with the first table
    source_table = source_doc.tables[0]
    target_table = target_doc.tables[0]
    # Iterate through each cell in the source table to copy formatting
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            # Copy font size and cell shading from source to target, if applicable
            # This is a simplified example; actual implementation may require copying more properties
            if i < len(target_table.rows) and j < len(target_table.rows[i].cells):
                target_cell = target_table.rows[i].cells[j]
                # Example: Copying paragraph text and formatting
                # This is a basic example and might not cover all formatting aspects
              #  target_cell.text = cell.text
                # Copying font size and style
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if i < len(target_cell.paragraphs) and len(target_cell.paragraphs[i].runs) > 0:
                            target_run = target_cell.paragraphs[i].runs[0]
                            target_run.font.size = run.font.size
                            target_run.font.bold = run.font.bold
                            target_run.font.italic = run.font.italic

            else:
                 pass               
                # Additional formatting aspects (e.g., cell shading, borders) would require further code
    # Save the modified target document
    return target_doc

# Main Streamlit app
def read_word_file():
    dir = os.getcwd()
    file_path = os.path.join(dir, "template.docx")

    #print(file_path)
    #print("Current Working Directory:", os.getcwd())
    book_data = process_book_text(book_text)
    print('book_data',type(book_data))
    print('book_data',book_data)
    if file_path:
        try:
            # Read the Word document
            with open(file_path, 'rb') as doc_file:
                uploaded_file = Document(doc_file)

                # Display tables from the Word document
                #st.write("Read word file")

        except Exception as e:
            st.error(f"Error reading tables from the Word document: {e}")
            return

        if uploaded_file is not None:
           # try:
                # Provide variable data based on processed book data
                variable_data = [
                    book_data["book_title"],
                    book_data["author_name"],
                    book_data["book_subtitle"],
                    book_data["tagline_suggestions"],
                    book_data["keywords"],
                    book_data["bisac_codes"],
                    book_data["online_book_description"],
                    book_data["long_back_cover_copy"],
                    book_data["short_back_cover_copy"],
                    book_data["about_the_author"],
                    book_data["headline"]
                ]

                # Read and update tables from the Word document
                updated_document = read_and_update_tables(uploaded_file, variable_data)

                # Save the updated document
                updated_document_bytes = BytesIO()
                updated_document.save(updated_document_bytes)
                #st.subheader("Updated Tables in Word Document:")
                #for i, updated_table in enumerate(updated_document.tables, start=1):
                    #st.write(f"Table {i}:")
                    #st.table([[cell.text for cell in row.cells] for row in updated_table.rows])

            # Create a download link for the updated Word document
                st.markdown(get_binary_file_downloader_html(updated_document_bytes.getvalue(), filename=file_name_input + "-output"), unsafe_allow_html=True)
                #get_binary_file_downloader_html(updated_document_bytes.getvalue(), filename="generated_text")
                #st.markdown(get_binary_file_downloader_html_new(updated_document_bytes.getvalue(), filename="generated_text"), unsafe_allow_html=True)
                # Automatically trigger the download using JavaScript
                #get_binary_file_downloader_html(updated_document_bytes.getvalue(), filename="generated_text")
           # except Exception as e:
            #    st.error(f"Error reading or updating tables from the Word document: {e}")
                # Display the updated table content in the app



def get_binary_file_downloader_html(bin_data, filename, file_extension=".docx"):
    b64 = base64.b64encode(bin_data).decode()
    return f'<a href="data:application/msword;base64,{b64}" download="{filename}{file_extension}">Click to download</a>'
    #download_link = f'<a id="download_link" href="data:application/msword;base64,{b64}" download="{filename}{file_extension}">Click to download</a>'

    # Display the link
    #st.markdown(download_link, unsafe_allow_html=True)

    # Automatically trigger the download using JavaScript
  




# Function to save the generated text as a Word document
def save_as_word_document(result, filename):
    doc = Document()

    # Add a table with two columns
    table = doc.add_table(rows=1, cols=2)

    # Set column names in the first row
    table.cell(0, 0).text = "Column Name"
    table.cell(0, 1).text = "Generated Text"

    # Add the result to the table
    row = table.add_row()
    row.cells[0].text = "Your Column Name"
    row.cells[1].text = result

    # Save the document
    doc.save(filename)

# Function to extract and convert tables to DataFrames
def extract_table_to_dataframe(table):
    data = []

    # Extract the first row of the table as column names
    columns = [cell.text for cell in table.rows[0].cells]

    # Check for and handle duplicate column names
    seen = set()
    unique_columns = []
    for column in columns:
        if column in seen:
            i = 1
            new_column = f"{column}_{i}"
            while new_column in seen:
                i += 1
                new_column = f"{column}_{i}"
            unique_columns.append(new_column)
            seen.add(new_column)
        else:
            unique_columns.append(column)
            seen.add(column)

    # Iterate through the remaining rows
    data = []
    for row in table.rows[1:]:
        row_data = [cell.text for cell in row.cells]
        data.append(row_data)

    df = pd.DataFrame(data, columns=unique_columns)

    return df

def generate_content(prompt):

    formatted_prompt = prompt.format(
        Description=Description,
        Audience=Audience,
        BookTitle=BookTitle,
        AuthorName=AuthorName,
        SubTitle = SubTitle,
        Genre = Genre,
        OnlineBookDescription= OnlineBookDescription,
        LongBackCoverCopy = LongBackCoverCopy,
        Headline = Headline,
        AuthorBiography=AuthorBiography,
        KeywordsForOnlineDiscoverability=KeywordsForOnlineDiscoverability
    )
   
    
    base_prompt='You are an expert in book publishing and Marketing. Look at the process below and answer like you are an Expert'
    complete_prompt=base_prompt+'\n'+formatted_prompt
    llm = AzureChatOpenAI(
    openai_api_version="2024-02-15-preview",
    azure_deployment="gpt-4",
    openai_api_key=azure_api_key,
    azure_endpoint = azure_endpoint
)
    
    query_turbo_1 = llm.invoke(complete_prompt)

    #st.write(query_turbo_1)
    return query_turbo_1.content
    
def generate_results_complete():
    # Define the prompt with placeholders
    prompt = """
    The following is a brief overview of what the book is about {Description}. Below are separate sets of instructions please create separate outputs on one document following the instructions below step by step. Make sure each step is carried out completely, do NOT skip any piece of instructions I have given you.

    Book Title:
    [Instructions]:
    Place the book's title {BookTitle} here. Do not make this up, only copy and paste from the provided document.

    Author Name: 
    [Instructions]
    Place the author's name {AuthorName} here. Do not make this up, only copy and paste from the provided document.

    Book Subtitle:
    [Instructions]
    Place the book's subtitle {SubTitle} here. Do not make this up, only copy and paste from the provided document.

    Tagline Suggestions:
    [Instructions]
    Create a section made up of 2-5  tagline suggestions for this book. The taglines should not include the book title or author name.

    Keywords for Online Discoverability:
    [Instructions]
    Create 15 Keywords for Online Discoverability. These keywords are related to the book and will help with online discoverability. Each word should be separated by semicolons.

    BISAC Codes:
    [Instructions]
    Provide only 3 BISAC Codes relating to the {Genre} of the book that are as specific as possible. NEVER make up the BISAC Codes. ONLY use accurate and real codes that relate to the {Genre} in the output. Keep this in a list format separated by semicolons. For example this section should look exactly like “JUV039140 JUVENILE FICTION ; Social Themes ; Friendship JUV039060 JUVENILE FICTION ; Social Themes ; Self-Esteem & Self-Reliance JUV002030 JUVENILE FICTION ; Animals ; Frogs & Toads” but the codes and words will change and be relevant to the book {Genre}. NEVER make up the BISAC Codes. ONLY use accurate and real codes in the output.


    Online Book Description:
    [Instructions]
    Write a 300 word Online Book Description in paragraph format.
    The Online Book Description will be optimized for Amazon. This is a book description that will be seen on the retailer websites where the book is available for sale.
    Focus on marketability and advertising to the audience of {Audience}. This will ALWAYS include the reading level at which this section should be written.

    Long Back Cover Copy:
    [Instructions]
    Write a 250 word long back cover summary using {Description} for context on the book.
        -Format for back cover summary:
        -The word count is 250 words every time. 
        -Focus on summarizing the plot in an exciting way for the {Audience} in 250 words. 

    Short Back Cover Copy:
    [Instructions]
    Write a 75 word Short back cover summary of the book based on {Description}. This summary will be different from the previous two. The goal is summarizing the plot to make the book desirable to a reader. Make sure each step is carried out completely, do NOT skip or stray from any instructions I have given you.

    About the Author:
    [Instructions]
    Write a 100 word author biography that gives information about the author. Find this information in the document provided to you and do not make information about the author up. The book is called {BookTitle} and the author is {AuthorName}. Keep the author biography’s language in line with language and reading levels that would be appropriate for {Audience}. If there is more than one author, create multiple but separate biographies. Each author on a project should have their own biography/ description. All will follow the instructions above and the style guide below.

    Headline:
    [Instructions]
    Write a headline of 239 characters including spaces and punctuation. 
    Write a headline for the book provided.
    Do not skip this step.
    The headline should not include the author name or title. 
    This is the elevator pitch of the book.

    Note:
    [Instructions]
    Copy and paste the following note: “Note: The given BISAC codes are based on the descriptions available; you might want to double-check them when you actually use them for publishing purposes.”

    STYLE GUIDE:
    NEVER include quotation marks or brackets in the output. Only provide what is specifically asked of you.
    Never put “[BOLD TITLE]” or “[SECTION TEXT]” or “[SECTION TITLE] in any output ever. 
    All content should be written at a reading level that is appropriate for the {Audience}. Age plays a factor in what a group of readers will be able to understand. Make sure all output ALWAYS stays at the reading level of the {Audience}.
    All sections will ALWAYS follow the word count guidelines EXACTLY without exception. 
    Word counts are REQUIREMENTS not suggestions. Be strict in following the word or character counts for EVERY section above. Any numbers included in the prompt are the word count requirements. For example “write a 300 word online book description…” the number 300 is the word count. This is accurate for every section and number provided to you. 
    Keep all text aligned nicely.
    Do not include any language from the prompt in any output.
    Make sure each step is carried out completely, do NOT skip steps or stray from ANY instructions you have been given.
    """
    # Use string formatting to replace placeholders with actual values
    formatted_prompt = prompt.format(
        Description=Description,
        Audience=Audience,
        BookTitle=BookTitle,
        AuthorName=AuthorName,
        SubTitle = SubTitle,
        AuthorBiography=AuthorBiography,
        Genre = Genre
        
    )

    # Display the formatted prompt
    #st.write(formatted_prompt)

    # Code to run the open ai for generation 
    st.subheader("Output Generated")
    base_prompt='You are an expert in book publishing and Marketing. Look at the process below and answer like you are an Expert'
    complete_prompt=base_prompt+'\n'+formatted_prompt
    llm = AzureChatOpenAI(
    openai_api_version="2024-02-15-preview",
    azure_deployment="gpt-4",
    openai_api_key=azure_api_key,
    azure_endpoint = azure_endpoint
)
    
    query_turbo_1 = llm.invoke(complete_prompt)

    #st.write(query_turbo_1)
    st.write(query_turbo_1.content)
    book_text =  query_turbo_1.content
    if st.button("Save Generated Text as Word Document"):
        read_word_file()
        st.subheader("Prompt")
        st.write(prompt)

    # Create a Word document with the generated text
       # word_doc = BytesIO()
        #save_as_word_document(query_turbo_1, word_doc)
        

#########################################################################################################        
# Define your username and password
USERNAME = "Survey"
PASSWORD = "Amazon"

# Function to check if the provided credentials match
def check_credentials(username, password):
    return username == USERNAME and password == PASSWORD

def login_page():
    #st.sidebar.title("Login")
    st.title("Login")
    st.write("Username and Password are case sensitive")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if check_credentials(username, password):
            # Correct credentials; set the session state to indicate logged in
            st.session_state["logged_in"] = True
           # st.experimental_rerun()
        else:
            # Incorrect credentials; show an error message
            st.error("Incorrect username or password.")

# Check if the 'logged_in' key exists in the session state and if it's True
if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False  # Initialize it as False

if st.session_state["logged_in"]:

    if st.button('Logout'):
        st.session_state["logged_in"] = False
        # Clear the username and password fields in the sidebar
        if 'username' in st.session_state: del st.session_state['username']
        if 'password' in st.session_state: del st.session_state['password']
        st.experimental_rerun()  # Rerun the app to reflect the logout state

    st.title("PalmettoPublishing")   #Main function starts here
    st.subheader("29-Feb")

    # Create a sidebar

    st.sidebar.title("Menu")
    #uploaded_file = st.file_uploader("Upload a DOCX File", type=["docx"])
    uploaded_file = st.sidebar.file_uploader("Upload a DOCX or DOC File", type=["docx", "doc"])

    combined_df = pd.DataFrame()
    if uploaded_file is not None:
        file_name_input = uploaded_file.name
        print('file_name_input',file_name_input)
        file_name_input, file_extension = os.path.splitext(file_name_input)
        target_docx='reference.docx'

        if file_extension.lower() == '.doc':
            # Convert DOC to DOCX 
            import aspose.words as aw
            print('uploaded_file',uploaded_file)
            temp_dir = tempfile.mkdtemp()
            path = os.path.join(temp_dir, uploaded_file.name)
            with open(path, "wb") as f:
                    f.write(uploaded_file.getvalue())
            doc_pr = aw.Document(path)
            converted_path=os.path.splitext(path)[0] + '.docx'
            doc_pr.save(converted_path)
            # doc=copy_table_format(converted_path, target_docx)
        #  doc=Document(converted_path)
            doc = dct()
                # Load a Word document
            doc.LoadFromFile(converted_path)

        elif file_extension.lower() == '.docx':
            # For DOCX files
        #   doc=copy_table_format(uploaded_file, target_docx)
        #  doc = Document(uploaded_file)
            temp_dir = tempfile.mkdtemp()
            path = os.path.join(temp_dir, uploaded_file.name)
            with open(path, "wb") as f:
                    f.write(uploaded_file.getvalue())
            doc = dct()
            doc.LoadFromFile(path)

        else:
            print("Unsupported file format. Please upload a .doc or .docx file.")
        

        df=extract_tabular_data(doc)
        combined_df = pd.concat([combined_df, df], axis=1)

        # Fill NaN values with empty strings
        combined_df = combined_df.fillna('')

        # Get the column names
        column_names = combined_df.columns.tolist()

        st.subheader("Survey data read from the file")
        st.dataframe(combined_df)
    
        # Labels to search for
        # Labels to search for
        labels_to_search = ['Author Name', 'summary', 'target audience', 'Book title', 'specific genre', 'Subtitle', 'Author Biography', 'Author Bio']

    # Dictionary to store the row numbers for each label
        label_rows = {}
        # Initialize variables with default values
        variables = {
            'Author Name': None,
            'summary': None,
            'target audience': None,
            'Book title': None,
            'specific genre': None,
            'Subtitle': None
        }
    # Loop through the labels and find their row numbers
        for label in labels_to_search:
            try:
                # Find the row index where the label keyword is present in the first column
                row_index = combined_df[combined_df.iloc[:, 0].str.contains(label, case=False)].index[0]
                # Store the row number in the dictionary
                label_rows[label] = row_index
                # Assign values to variables using the dictionary
                variables[label] = combined_df.iloc[row_index, 1]
            except IndexError:
                if label=='Author Biography':
                    st.warning(f'Warning: Label "{label}" is missing checking for Author Bio')
                else:
                    # Handle the case where the label is not found in the DataFrame
                    st.warning(f'Warning: Label "{label}" is missing and can impact ouput results')
                    
            
    # Print the results
        #for label, row in label_rows.items():
        #   st.write(f"{label}: {combined_df.iloc[row, 0]}")    
        AuthorName = variables['Author Name']
        Description = variables['summary']
        Audience = variables['target audience']
        BookTitle = variables['Book title']
        Genre = variables['specific genre']
        Subtitle = variables['Subtitle']
        # Handle Author Biography and Author Bio
        if 'Author Biography' in label_rows and 'Author Bio' in label_rows:
            # If both labels are found, choose one (you can modify this logic as needed)
            AuthorBiography = combined_df.iloc[label_rows['Author Biography'], 1]
        elif 'Author Biography' in label_rows or 'Author Bio' in label_rows:
    # Use the label that is found
             AuthorBiography = combined_df.iloc[label_rows.get('Author Biography', label_rows.get('Author Bio')), 1]
        else:
    # Both labels are missing, assign None
            AuthorBiography = None
          


        
        Tagline_Suggestions_prompt = st.sidebar.text_area("Tagline Suggestions:", value=Tagline_Suggestions_prompt, height=100)
        Keywords_Online_Discoverability_prompt = st.sidebar.text_area("Keywords for Online Discoverability:", value=Keywords_Online_Discoverability_prompt, height=100)
        BISAC_Codes_prompt = st.sidebar.text_area("BISAC Codes:", value=BISAC_Codes_prompt, height=100)
        Headline_prompt = st.sidebar.text_area("Headline:", value=Headline_prompt, height=100)
        Online_Book_Description_prompt = st.sidebar.text_area("Online Book Description:", value=Online_Book_Description_prompt, height=100)
        Long_Back_Cover_Copy_prompt = st.sidebar.text_area("Long Back Cover Copy:", value=Long_Back_Cover_Copy_prompt, height=100)
        Short_Back_Cover_Copy_prompt = st.sidebar.text_area("Short Back Cover Copy:", value=Short_Back_Cover_Copy_prompt, height=100)
        about_the_author_prompt = st.sidebar.text_area("About the Author:", value=about_the_author_prompt, height=100)


        submit_button = st.sidebar.button('Submit Button')

    # Initialize session state
        if 'all_generated_content' not in st.session_state:
            st.session_state.all_generated_content = ""

        # Your existing code
        if submit_button:
            # Reset the content if the button is clicked
            st.session_state.all_generated_content = ""
            
            st.session_state.screenOutput=""
            # Add existing st.write statements to the variable and display using st.write
            st.session_state.all_generated_content += "Book Title: " + BookTitle + "\n"
            st.write("Book Title: " + BookTitle)
            st.session_state.screenOutput += "Book Title: " + BookTitle + "\n"


            st.session_state.all_generated_content += "\n" + "Author Name: " + AuthorName + "\n"
            st.write("Author Name: " + AuthorName)
            
            st.session_state.screenOutput += "\n" + "Author Name: " + AuthorName + "\n"

            st.session_state.all_generated_content += "\n" + "Book Subtitle: " + Subtitle + "\n"
            st.write("Book Subtitle: " + Subtitle)
            st.session_state.screenOutput += "\n" + "Book Subtitle: " + Subtitle + "\n" 


            # Add generated content to the variable and display using st.write
            TaglineSuggestion_gen = generate_content(Tagline_Suggestions_prompt)
            st.session_state.all_generated_content += "\n" + "Tagline Suggestions: " + TaglineSuggestion_gen + "\n"
            st.write("Tagline Suggestions: " + TaglineSuggestion_gen)
            st.session_state.screenOutput  += "\n" + "Tagline Suggestions: " + TaglineSuggestion_gen + "\n"

            Keywords_gen = generate_content(Keywords_Online_Discoverability_prompt)
            st.session_state.all_generated_content += "\n" + "Keywords for Online Discoverability: " + Keywords_gen + "\n"
            st.write("Keywords for Online Discoverability: " + Keywords_gen)
            st.session_state.screenOutput  += "\n" + "Keywords for Online Discoverability: " + Keywords_gen + "\n"
            KeywordsForOnlineDiscoverability = Keywords_gen

            BISAC_gen = generate_content(BISAC_Codes_prompt)
            st.session_state.all_generated_content += "\n" + "BISAC Codes: " + BISAC_gen + "\n"
            st.write("BISAC Codes: " + BISAC_gen)
            st.session_state.screenOutput  += "\n" + "BISAC Codes: " + BISAC_gen + "\n"


            Headline_gen = generate_content(Headline_prompt)
            st.write("Headline: " + Headline_gen)
            Headline = Headline_gen
            st.session_state.head_match = Headline
            st.session_state.screenOutput   += "\n" + "Headline: " + Headline_gen + "\n"

            Online_gen = generate_content(Online_Book_Description_prompt)
            st.session_state.all_generated_content += "\n" + "Online Book Description: " + Online_gen + "\n"
            st.write("Online Book Description: " + Online_gen)
            OnlineBookDescription = Online_gen
            st.session_state.screenOutput  += "\n" + "Online Book Description: " + Online_gen + "\n"

            Back_cover_gen = generate_content(Long_Back_Cover_Copy_prompt)
            st.session_state.all_generated_content += "\n" + "Long Back Cover Copy: " + Back_cover_gen + "\n"
            st.write("Long Back Cover Copy: " + Back_cover_gen)
            LongBackCoverCopy= Back_cover_gen
            st.session_state.screenOutput   += "\n" + "Long Back Cover Copy: " + Back_cover_gen + "\n"

            Short_cover_gen = generate_content(Short_Back_Cover_Copy_prompt)
            st.session_state.all_generated_content += "\n" + "Short Back Cover Copy: " + Short_cover_gen + "\n"
            st.write("Short Back Cover Copy: " + Short_cover_gen)
            st.session_state.screenOutput   += "\n" + "Short Back Cover Copy: " + Short_cover_gen + "\n"


            About_author_gen = generate_content(about_the_author_prompt)
            st.session_state.all_generated_content += "\n" + "About the Author: " + About_author_gen + "\n"
            st.write("About the Author: " + About_author_gen)
            st.session_state.all_generated_content += "\n" + "Headline: " + Headline_gen + "\n"
            st.session_state.screenOutput   += "\n" + "About the Author: " + About_author_gen + "\n"

    # Now, display all generated content using a single st.write statement with newlines

            # Now, display all generated content using a single st.write statement with newlines
            #st.write("Final Results")
            #st.write(all_generated_content)
        #st.write("All generated Content")
        #st.write(st.session_state.all_generated_content)

        
        book_text =  st.session_state.all_generated_content
        if st.sidebar.button("Save Output as Word Document"):
            st.write(st.session_state.screenOutput)   
            read_word_file()
else:
    print("Please login")
    login_page()  # Show the login page if not logged in
